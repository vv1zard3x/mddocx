"""Markdown -> DOCX conversion engine.

The high-level entry point is :func:`convert`. It expects a folder
containing exactly one ``*.md`` file and optionally a ``media/``
subdirectory referenced by image links. Output is a single ``.docx``
formatted according to the GOST 7.32 rules expressed in
:mod:`docxmd._styles`.

See ``MARKDOWN_RULES.md`` for the markdown contract the converter
relies on. Anything outside that contract is best-effort.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as _Doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image as _DocxImage
from docx.opc.constants import RELATIONSHIP_TYPE as _RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt
from docx.text.paragraph import Paragraph
from markdown_it import MarkdownIt
from markdown_it.token import Token

from ._numbering import (
    BULLET_NUM_ID,
    ListContext,
    allocate_ordered_num_id,
    apply_numbering,
    install_numbering,
)
from ._styles import (
    CODE_SHADING_HEX,
    FIRST_LINE_INDENT_CM,
    LINK_COLOR_HEX,
    MAX_CONTENT_WIDTH_CM,
    MONO,
    apply_gost_styles,
    set_paragraph_horizontal_rule,
    set_run_font,
    set_run_shading,
)

# Per GOST 7.32 captions use the EM DASH separator: "Рисунок 1 — Название".
_CAPTION_SEPARATOR = "\u00a0\u2014\u00a0"  # NBSP + em dash + NBSP
_TABLE_CAPTION_PREFIX_RE = re.compile(
    r"^\s*табл(?:ица)?\s*[\.:\u2014\-]\s*",
    re.IGNORECASE,
)


class ConversionError(RuntimeError):
    """Raised when the source folder layout is invalid."""


def convert(src: Path | str, output: Path | str | None = None) -> Path:
    """Convert *src* directory (md + optional media/) into a .docx file.

    Parameters
    ----------
    src:
        Directory containing exactly one ``*.md`` file. Image paths in
        the markdown must be relative to this directory.
    output:
        Output ``.docx`` path. Defaults to ``<md filename>.docx`` next to
        the source markdown.

    Returns
    -------
    Path to the written ``.docx``.
    """

    src_dir = Path(src).expanduser().resolve()
    if not src_dir.is_dir():
        raise ConversionError(f"Source is not a directory: {src_dir}")

    md_files = sorted(p for p in src_dir.iterdir() if p.suffix.lower() == ".md")
    if not md_files:
        raise ConversionError(f"No .md file found in {src_dir}")
    if len(md_files) > 1:
        names = ", ".join(p.name for p in md_files)
        raise ConversionError(
            f"Expected exactly one .md file in {src_dir}, found: {names}"
        )
    md_path = md_files[0]

    if output is None:
        out_path = md_path.with_suffix(".docx")
    else:
        out_path = Path(output).expanduser().resolve()
        if out_path.is_dir():
            out_path = out_path / md_path.with_suffix(".docx").name

    text = md_path.read_text(encoding="utf-8")

    doc = Document()
    apply_gost_styles(doc)
    install_numbering(doc)
    _strip_initial_empty_paragraph(doc)

    parser = (
        MarkdownIt("commonmark", {"breaks": False, "html": False})
        .enable("table")
        .enable("strikethrough")
    )
    tokens = parser.parse(text)

    builder = DocBuilder(doc=doc, src_dir=src_dir)
    builder.render(tokens)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def _strip_initial_empty_paragraph(doc: _Doc) -> None:
    """Remove the blank paragraph python-docx inserts in fresh documents.

    Avoids a leading empty line in the rendered output without leaving
    the body element empty (Word rejects bodies with no block content).
    Re-added implicitly when the first block is appended.
    """

    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for p in list(body.findall(qn("w:p"))):
        if not p.text and len(p.findall(qn("w:r"))) == 0:
            body.remove(p)
    if sectPr is None:
        return
    # body must have at least one block-level element before sectPr; add a
    # placeholder paragraph if we removed everything else. python-docx will
    # happily append new blocks before sectPr afterwards.
    if not body.findall(qn("w:p")) and not body.findall(qn("w:tbl")):
        placeholder = OxmlElement("w:p")
        body.insert(list(body).index(sectPr), placeholder)


# --------------------------------------------------------------------------- #
# Builder
# --------------------------------------------------------------------------- #


@dataclass
class DocBuilder:
    doc: _Doc
    src_dir: Path
    image_counter: int = 0
    table_counter: int = 0
    _list_stack: list[ListContext] = field(default_factory=list)
    _quote_depth: int = 0
    # Pending caption text and its original inline tokens, kept while we peek
    # ahead one block to see whether a table follows.
    _pending_table_caption: tuple[str, Sequence[Token]] | None = None

    # ----- public entry ----- #

    def render(self, tokens: Sequence[Token]) -> None:
        i = 0
        n = len(tokens)
        while i < n:
            t = tokens[i]
            tt = t.type

            # Flush a pending table caption if the next block is not a table.
            if (
                self._pending_table_caption is not None
                and tt != "table_open"
            ):
                _, fallback_inline = self._pending_table_caption
                self._pending_table_caption = None
                self._add_paragraph(fallback_inline)

            if tt == "heading_open":
                level = int(t.tag[1:])
                inline = tokens[i + 1]
                self._add_heading(level, inline.children or [])
                i += 3  # heading_open, inline, heading_close
                continue

            if tt == "paragraph_open":
                inline = tokens[i + 1]
                children = inline.children or []
                # Detect "Таблица: <name>" paragraphs at the top level (not
                # inside a list/quote) — they become the next table's caption.
                if not self._list_stack and not self._quote_depth:
                    plain = _collect_plain_text(children).strip()
                    m = _TABLE_CAPTION_PREFIX_RE.match(plain)
                    if m:
                        title = plain[m.end():].strip().rstrip(".")
                        self._pending_table_caption = (title, children)
                        i += 3
                        continue
                self._add_paragraph(children)
                i += 3
                continue

            if tt == "bullet_list_open":
                self._push_list("bullet")
                i += 1
                continue
            if tt == "ordered_list_open":
                self._push_list("ordered")
                i += 1
                continue
            if tt in ("bullet_list_close", "ordered_list_close"):
                if self._list_stack:
                    self._list_stack.pop()
                i += 1
                continue

            if tt == "list_item_open" or tt == "list_item_close":
                i += 1
                continue

            if tt == "blockquote_open":
                self._quote_depth += 1
                i += 1
                continue
            if tt == "blockquote_close":
                if self._quote_depth:
                    self._quote_depth -= 1
                i += 1
                continue

            if tt in ("fence", "code_block"):
                self._add_code_block(t.content.rstrip("\n"))
                i += 1
                continue

            if tt == "hr":
                self._add_hr()
                i += 1
                continue

            if tt == "table_open":
                if self._pending_table_caption is not None:
                    title, _ = self._pending_table_caption
                    self._pending_table_caption = None
                    self._add_table_caption(title)
                i = self._add_table(tokens, i)
                continue

            if tt == "html_block":
                # Skip raw HTML blocks; they fall outside the contract.
                i += 1
                continue

            i += 1

        # End of document: a stray pending caption falls back to plain text.
        if self._pending_table_caption is not None:
            _, fallback_inline = self._pending_table_caption
            self._pending_table_caption = None
            self._add_paragraph(fallback_inline)

    # ----- list bookkeeping ----- #

    def _push_list(self, kind: str) -> None:
        """Update the active-list stack, allocating numIds as needed.

        A nested list of the same kind continues the parent's hierarchy
        (same ``numId``, deeper ``ilvl``); switching kind (or starting a
        top-level list) opens a fresh counter.
        """

        if self._list_stack and self._list_stack[-1].kind == kind:
            parent = self._list_stack[-1]
            self._list_stack.append(
                ListContext(kind=kind, num_id=parent.num_id, ilvl=parent.ilvl + 1)
            )
            return
        if kind == "bullet":
            num_id = BULLET_NUM_ID
        else:
            num_id = allocate_ordered_num_id(self.doc)
        self._list_stack.append(ListContext(kind=kind, num_id=num_id, ilvl=0))

    # ----- block builders ----- #

    def _add_heading(self, level: int, inline_tokens: Sequence[Token]) -> None:
        # H1 keeps its centered style (no indent); H2..H6 inherit 1.25 cm
        # first-line indent and pStyle-bound numbering from the heading
        # style configured in ``_styles.apply_gost_styles``. Touching
        # ``first_line_indent`` directly here would override the
        # abstractNum's <w:ind>, breaking auto-number placement — let the
        # style/numbering chain do its job.
        level = max(1, min(level, 9))
        p = self.doc.add_paragraph(style=f"Heading {level}")
        self._render_inline(p, inline_tokens)

    def _add_paragraph(self, inline_tokens: Sequence[Token]) -> None:
        if (
            not self._list_stack
            and not self._quote_depth
            and _is_standalone_image(inline_tokens)
        ):
            self._add_image_with_caption(_find_image(inline_tokens))
            return

        if self._list_stack:
            ctx = self._list_stack[-1]
            p = self.doc.add_paragraph()
            apply_numbering(p, num_id=ctx.num_id, ilvl=ctx.ilvl)
            # NB: do NOT touch first_line_indent / left_indent on list paras.
            # Per OOXML resolution order numbering > style, so the level's
            # <w:ind> from the abstractNum already overrides Normal's
            # firstLine=1.25cm. Writing any direct <w:ind> here (even
            # firstLine=0) causes renderers to replace the numbering's ind
            # wholesale — newly typed items in Word then visibly use a
            # different indent than the converted ones (the very symptom of
            # "положение 1.1. совпадает с 1." reported during manual review).
            # Also override Normal's "justify": list items read better
            # ragged-right when the marker creates a visible hanging.
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self._quote_depth:
            p = self.doc.add_paragraph(style="Quote")
            indent_cm = 1.25 * self._quote_depth
            p.paragraph_format.left_indent = Cm(indent_cm)
        else:
            p = self.doc.add_paragraph()

        self._render_inline(p, inline_tokens)

    def _add_table_caption(self, title: str) -> None:
        self.table_counter += 1
        p = self.doc.add_paragraph(style="Table Caption")
        p.paragraph_format.first_line_indent = Cm(0)
        if title:
            text = f"Таблица {self.table_counter}{_CAPTION_SEPARATOR}{title}"
        else:
            text = f"Таблица {self.table_counter}"
        p.add_run(text)

    def _add_code_block(self, content: str) -> None:
        for line in content.split("\n"):
            p = self.doc.add_paragraph(style="Code Block")
            p.paragraph_format.first_line_indent = Cm(0)
            if line:
                run = p.add_run(line)
                set_run_font(run, MONO)
            # empty line still produces a paragraph in the block

    def _add_hr(self) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        set_paragraph_horizontal_rule(p)

    def _add_table(self, tokens: Sequence[Token], i: int) -> int:
        rows: list[tuple[bool, list[list[Token]], list[str | None]]] = []
        current_row_cells: list[list[Token]] | None = None
        current_row_aligns: list[str | None] | None = None
        current_cell_inline: list[Token] | None = None
        current_cell_align: str | None = None
        is_header = False
        i += 1
        while tokens[i].type != "table_close":
            t = tokens[i]
            tt = t.type
            if tt == "thead_open":
                is_header = True
            elif tt == "thead_close":
                is_header = False
            elif tt == "tr_open":
                current_row_cells = []
                current_row_aligns = []
            elif tt == "tr_close":
                if current_row_cells is not None and current_row_aligns is not None:
                    rows.append((is_header, current_row_cells, current_row_aligns))
                current_row_cells = None
                current_row_aligns = None
            elif tt in ("th_open", "td_open"):
                current_cell_inline = []
                current_cell_align = t.attrs.get("style") if t.attrs else None
            elif tt in ("th_close", "td_close"):
                if current_row_cells is not None and current_cell_inline is not None:
                    current_row_cells.append(current_cell_inline)
                    current_row_aligns.append(current_cell_align)  # type: ignore[union-attr]
                current_cell_inline = None
                current_cell_align = None
            elif tt == "inline":
                current_cell_inline = list(t.children or [])
            i += 1

        if not rows:
            return i + 1

        n_cols = max(len(cells) for _, cells, _ in rows)
        table = self.doc.add_table(rows=len(rows), cols=n_cols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.autofit = True

        for ridx, (header, cells, aligns) in enumerate(rows):
            for cidx in range(n_cols):
                cell = table.rows[ridx].cells[cidx]
                cell_inline = cells[cidx] if cidx < len(cells) else []
                align_style = aligns[cidx] if cidx < len(aligns) else None
                p = cell.paragraphs[0]
                p.text = ""
                p.paragraph_format.first_line_indent = Cm(0)
                if align_style:
                    if "center" in align_style:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif "right" in align_style:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._render_inline(p, cell_inline)
                if header:
                    for run in p.runs:
                        run.bold = True
        # blank paragraph after a table prevents the next paragraph from
        # being absorbed into the trailing cell when opened in Word.
        self.doc.add_paragraph()
        return i + 1

    def _add_image_with_caption(self, image_token: Token) -> None:
        src = image_token.attrs.get("src", "") if image_token.attrs else ""
        alt = (image_token.content or "").strip()
        if not alt and image_token.attrs:
            alt = (image_token.attrs.get("title") or "").strip()

        img_path = (self.src_dir / src).resolve() if src else None
        if not img_path or not img_path.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run(f"[Не найдено изображение: {src}]").italic = True
            return

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        width = _picture_width(img_path)
        run.add_picture(str(img_path), width=width)

        self.image_counter += 1
        caption_text = (
            f"Рисунок {self.image_counter}{_CAPTION_SEPARATOR}{alt}"
            if alt
            else f"Рисунок {self.image_counter}"
        )
        cap = self.doc.add_paragraph(style="Caption")
        cap.add_run(caption_text)

    # ----- inline rendering ----- #

    def _render_inline(self, paragraph: Paragraph, tokens: Sequence[Token]) -> None:
        bold = 0
        italic = 0
        strike = 0
        i = 0
        n = len(tokens)
        while i < n:
            t = tokens[i]
            tt = t.type
            if tt == "text":
                self._add_text_run(paragraph, t.content, bold, italic, strike)
            elif tt == "softbreak":
                self._add_text_run(paragraph, " ", bold, italic, strike)
            elif tt == "hardbreak":
                paragraph.add_run().add_break()
            elif tt == "strong_open":
                bold += 1
            elif tt == "strong_close":
                bold = max(0, bold - 1)
            elif tt == "em_open":
                italic += 1
            elif tt == "em_close":
                italic = max(0, italic - 1)
            elif tt == "s_open":
                strike += 1
            elif tt == "s_close":
                strike = max(0, strike - 1)
            elif tt == "code_inline":
                self._add_code_run(paragraph, t.content)
            elif tt == "link_open":
                href = (t.attrs or {}).get("href", "") if t.attrs else ""
                j = i + 1
                text_parts: list[str] = []
                while j < n and tokens[j].type != "link_close":
                    inner = tokens[j]
                    if inner.type == "text":
                        text_parts.append(inner.content)
                    elif inner.type == "softbreak":
                        text_parts.append(" ")
                    elif inner.type == "code_inline":
                        text_parts.append(inner.content)
                    j += 1
                link_text = "".join(text_parts) or href
                self._add_hyperlink(
                    paragraph, href, link_text, bold=bold, italic=italic
                )
                i = j  # at link_close; loop will i+=1
            elif tt == "image":
                src = (t.attrs or {}).get("src", "")
                alt = (t.content or "").strip()
                self._add_inline_image(paragraph, src, alt)
            elif tt == "html_inline":
                pass
            i += 1

    def _add_text_run(
        self,
        paragraph: Paragraph,
        text: str,
        bold: int,
        italic: int,
        strike: int,
    ) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True

    def _add_code_run(self, paragraph: Paragraph, text: str) -> None:
        run = paragraph.add_run(text)
        set_run_font(run, MONO)
        run.font.size = Pt(10.5)
        set_run_shading(run, CODE_SHADING_HEX)

    def _add_hyperlink(
        self,
        paragraph: Paragraph,
        url: str,
        text: str,
        *,
        bold: int = 0,
        italic: int = 0,
    ) -> None:
        if not url:
            self._add_text_run(paragraph, text, bold, italic, 0)
            return
        part = paragraph.part
        r_id = part.relate_to(url, _RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), LINK_COLOR_HEX)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(color)
        rPr.append(u)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        r.append(rPr)
        wt = OxmlElement("w:t")
        wt.text = text
        wt.set(qn("xml:space"), "preserve")
        r.append(wt)
        hyperlink.append(r)
        paragraph._p.append(hyperlink)

    def _add_inline_image(self, paragraph: Paragraph, src: str, alt: str) -> None:
        img_path = (self.src_dir / src).resolve() if src else None
        if not img_path or not img_path.exists():
            self._add_text_run(paragraph, f"[image: {src}]", 0, 1, 0)
            return
        run = paragraph.add_run()
        width = _picture_width(img_path)
        run.add_picture(str(img_path), width=width)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #


def _is_standalone_image(tokens: Iterable[Token]) -> bool:
    """True iff the inline children represent a single image with only whitespace.

    Mirrors how typical markdown editors render a lone ``![](...)`` line as a
    figure. Inline images sharing a paragraph with text remain inline.
    """

    saw_image = False
    for t in tokens:
        if t.type == "image":
            if saw_image:
                return False
            saw_image = True
        elif t.type == "text":
            if t.content.strip():
                return False
        elif t.type == "softbreak":
            continue
        else:
            return False
    return saw_image


def _collect_plain_text(tokens: Iterable[Token]) -> str:
    """Recover the visible text of an inline token sequence.

    Used to test paragraph contents against the table-caption convention
    without committing to a Word paragraph first.
    """

    parts: list[str] = []
    for t in tokens:
        if t.type == "text":
            parts.append(t.content)
        elif t.type in ("softbreak", "hardbreak"):
            parts.append(" ")
        elif t.type == "code_inline":
            parts.append(t.content)
    return "".join(parts)


def _find_image(tokens: Iterable[Token]) -> Token:
    for t in tokens:
        if t.type == "image":
            return t
    raise ValueError("No image token in inline sequence")


def _picture_width(path: Path):
    """Return an EMU width for ``run.add_picture`` that respects page width.

    Falls back to a conservative cap if image metadata is unreadable.
    """

    try:
        img = _DocxImage.from_file(str(path))
        dpi = img.horz_dpi or 96
        natural_cm = (img.px_width / dpi) * 2.54
    except Exception:
        return Cm(MAX_CONTENT_WIDTH_CM)
    if natural_cm > MAX_CONTENT_WIDTH_CM:
        return Cm(MAX_CONTENT_WIDTH_CM)
    return Cm(natural_cm)
